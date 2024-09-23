import streamlit as st
from langchain_google_genai import GoogleGenerativeAI
from langchain_core.output_parsers import JsonOutputParser
from langchain_core.document_loaders import BaseLoader
from langchain_core.documents import Document
from PyPDF2 import PdfReader
import docx
from pydantic import BaseModel, Field

# Template for generating questions and answers
template = """
Kindly help me to set {number} questions and corresponding answers based on this provided context

context: {context}
format_instructions: {format_instructions}
"""

# Function to format the document context
def format_context(docs):
    return "\n".join(doc.page_content for doc in docs)

# Pydantic model for question-answer format
class QA_parser(BaseModel):
    question: str = Field("Questions generated")
    answer: str = Field("Answer for each of the question")

# Custom loader class for different document formats
class LambdaStreamlitLoader(BaseLoader):
    def __init__(self, file) -> None:
        self.file = file

    def lazy_load(self):
        file_name = self.file.name
        *_, ext = file_name.split(".")

        if ext == "docx" or ext == "doc":
            doc = docx.Document(self.file)
            for paragraph in doc.paragraphs:
                for line in paragraph.text.split("\n"):
                    yield Document(page_content=line)

        elif ext == "pdf":
            doc = PdfReader(self.file)
            for page in doc.pages:
                for line in page.extract_text().split("\n"):
                    yield Document(page_content=line)

# Template for guiding the student on their answers
help_template = """
I have been asked this question: {question}
And then I provided this answer: {student_answer}
However, this is the actual answer: {ground_truth}
You are not to respond with the actual answer, just compare my answer to the right answer and guide me.
If my answer indicates that I don't know, guide me with a clearer step-by-step approach to make me understand it, with words of encouragement.
The answer doesn't have to be verbatim, tell me I am correct if I use other words that have the same meaning.
"""

# Initialize LLM for help evaluation
help_llm = GoogleGenerativeAI(model="gemini-pro")

# Initialize session state
if "qa_json" not in st.session_state:
    st.session_state.qa_json = None

if "count" not in st.session_state:
    st.session_state.count = 0

if "questions_generated" not in st.session_state:
    st.session_state.questions_generated = False

if "document" not in st.session_state:
    st.session_state.document = None

# Sidebar for file upload and question generation
with st.sidebar:
    file = st.file_uploader("Upload document here", type=["pdf", "docx", "doc"])

    if file:
        # Load the document only once when a new file is uploaded
        if st.session_state.document is None:
            loader = LambdaStreamlitLoader(file)
            st.session_state.document = list(loader.lazy_load())
            st.success("Document Loaded successfully")

        number_ip = st.number_input("How many questions would you like to set?", step=1, min_value=1, format="%i")

        if st.button("Generate") and st.session_state.document:
            with st.spinner("Generating questions..."):
                llm = GoogleGenerativeAI(model="gemini-pro")
                parser = JsonOutputParser()

                # Build the prompt and make the LLM call
                full_prompt = template.format(
                    context=format_context(st.session_state.document),
                    number=number_ip,
                    format_instructions=parser.get_format_instructions()
                )

                llm_response = llm(full_prompt)
                qas = parser.parse(llm_response)
                
                # Save the generated questions and answers in the session state
                st.session_state.qa_json = qas
                st.session_state.questions_generated = True
                st.success("Questions and answers generated successfully!")

# Main content area
st.title("Question and Answer Session")

if st.session_state.questions_generated and st.session_state.qa_json:
    qa_session = st.session_state.qa_json["questions"]
    
    try:
        current_qa = qa_session[st.session_state.count]
        current_q = current_qa["question"]
        current_a = current_qa["answer"]

        st.markdown(f"#### Question {st.session_state.count + 1}")
        st.write(f"Question: {current_q}")
        with st.expander("Reveal answer"):
            st.write("Actual answer from material:")
            st.write(current_a)


        # Ensure session state for the current question is only initialized once
        if f"user_answer_{st.session_state.count}" not in st.session_state:
            st.session_state[f"user_answer_{st.session_state.count}"] = ""

        # Render the text area with the value already stored in session state
        user_answer = st.text_area(
            "Provide your answer here", 
            value=st.session_state[f"user_answer_{st.session_state.count}"],
            key=f"user_answer_{st.session_state.count}"
        )

        # Store the input from the user dynamically
        if st.button("Evaluate with AI"):
            with st.spinner("Evaluating..."):
                help_message = help_template.format(
                    question=current_q, 
                    student_answer=user_answer, 
                    ground_truth=current_a
                )
                help_response = help_llm(help_message)
                st.write(help_response)

        if st.button("NEXT"):
            if user_answer:
                st.session_state.count += 1
                st.rerun()
            else:
                st.error("Please provide an answer. If unsure, you can mention that you don't know.")

    except IndexError:
        st.markdown("## Congratulations, you've completed the session! You may generate more questions to continue practicing.")
        st.session_state.count = 0
        st.session_state.questions_generated = False

    except Exception as e:
        st.error(f"An error occurred: {e}. Please regenerate the questions as the model may have been unstable.")
        
else:
    st.warning("Upload a document and specify the number of questions to proceed.")
